"""
Build DOCX template — single page, strict RTL, matched to reference image.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK_BLUE_HEX = "1F4E79"
LIGHT_BLUE_HEX = "D9EAF7"
SIG_BOX_MARKER = "⟦SIG_BOX⟧"
DATE_BOX_MARKER = "⟦DATE_BOX⟧"
ZEBRA_HEX = "F2F2F2"
WHITE_HEX = "FFFFFF"
TOTAL_ROW_HEX = "E8F0F8"
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)

FONT_NAME = "Arial"
BODY_SIZE = Pt(10.5)
TABLE_SIZE = Pt(10.5)
TITLE_SIZE = Pt(22)
SECTION_SIZE = Pt(11.5)
RECEIPT_TITLE_SIZE = Pt(12.5)
NOTES_SIZE = Pt(9.5)
FOOTER_SIZE = Pt(9.5)
TABLE_ROW_HEIGHT = 380
TABLE_HEADER_HEIGHT = 440


def _jc_val(align) -> str:
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if align == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    return "right"


def _set_paragraph_rtl(paragraph, *, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi", "w:textAlignment"):
        node = p_pr.find(qn(tag))
        if node is not None:
            p_pr.remove(node)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), _jc_val(align))
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)


def _set_run_rtl(run, *, bold=False, size=None, color=None, underline=False) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or BODY_SIZE
    run.font.underline = underline
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:rtl", "w:lang"):
        node = r_pr.find(qn(tag))
        if node is not None:
            r_pr.remove(node)
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    r_pr.append(OxmlElement("w:rtl"))
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    r_pr.append(lang)
    if color:
        run.font.color.rgb = color


def _set_run_ltr(run, *, bold=False, size=None, color=None, underline=False) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or BODY_SIZE
    run.font.underline = underline
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:rtl", "w:lang"):
        node = r_pr.find(qn(tag))
        if node is not None:
            r_pr.remove(node)
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:bidi"), "en-US")
    r_pr.append(lang)
    if color:
        run.font.color.rgb = color


def _para(doc, *, after=1, before=0, align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.15):
    p = doc.add_paragraph()
    _set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = line_spacing
    return p


def _text(doc, content, *, bold=False, color=None, size=None, after=1, before=0, align=WD_ALIGN_PARAGRAPH.RIGHT, underline=False):
    p = _para(doc, after=after, before=before, align=align)
    r = p.add_run(content)
    _set_run_rtl(r, bold=bold, color=color, size=size, underline=underline)
    return p


def _mixed(doc, parts, *, after=1):
    p = _para(doc, after=after)
    for content, style in parts:
        r = p.add_run(content)
        _set_run_rtl(r, bold=style.get("bold", False), color=style.get("color"), underline=style.get("underline", False))
    return p


def _shade(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _borders(cell, color="BFBFBF", dashed=False, sz="4") -> None:
    val = "dashed" if dashed else "single"
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), color)
        b.append(e)
    cell._tc.get_or_add_tcPr().append(b)


def _pad(cell, t=55, b=55, s=70, e=70) -> None:
    m = OxmlElement("w:tcMar")
    for name, v in (("top", t), ("bottom", b), ("start", s), ("end", e)):
        n = OxmlElement(f"w:{name}")
        n.set(qn("w:w"), str(v))
        n.set(qn("w:type"), "dxa")
        m.append(n)
    cell._tc.get_or_add_tcPr().append(m)


def _row_height(row, height: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for tag in ("w:trHeight",):
        node = tr_pr.find(qn(tag))
        if node is not None:
            tr_pr.remove(node)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(height))
    h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(h)


def _cell(
    cell,
    content,
    *,
    bold=False,
    fill=None,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    white=False,
    size=None,
    amount=False,
):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_rtl(p, align=align)
    text = content if not amount else f"\u200e{content}"
    r = p.add_run(text)
    if amount:
        _set_run_ltr(r, bold=bold, size=size or TABLE_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF) if white else None)
    elif white:
        _set_run_rtl(r, bold=bold, size=size or TABLE_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF))
    else:
        _set_run_rtl(r, bold=bold, size=size or TABLE_SIZE)
    _pad(cell, t=48 if not fill else 52, b=48 if not fill else 52)
    if fill:
        _shade(cell, fill)


def _tbl_pr(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    return tbl_pr


def _table_rtl(table) -> None:
    tbl_pr = _tbl_pr(table)
    for tag in ("w:bidiVisual", "w:jc", "w:tblInd"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    tbl_pr.append(OxmlElement("w:bidiVisual"))
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    tbl_pr.append(jc)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)


def _table_width(table, cm: float) -> None:
    tbl_pr = _tbl_pr(table)
    for tag in ("w:tblW", "w:tblLayout"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(cm * 567)))
    w.set(qn("w:type"), "dxa")
    tbl_pr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _set_doc_defaults_rtl(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)

    p_pr_default = doc_defaults.find(qn("w:pPrDefault"))
    if p_pr_default is not None:
        doc_defaults.remove(p_pr_default)
    p_pr_default = OxmlElement("w:pPrDefault")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))
    p_pr_default.append(p_pr)
    doc_defaults.append(p_pr_default)

    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is not None:
        doc_defaults.remove(r_pr_default)
    r_pr_default = OxmlElement("w:rPrDefault")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.append(r_fonts)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    r_pr.append(lang)
    r_pr.append(OxmlElement("w:rtl"))
    r_pr_default.append(r_pr)
    doc_defaults.append(r_pr_default)


def _set_settings_rtl(doc: Document) -> None:
    settings = doc.part.settings.element
    theme = settings.find(qn("w:themeFontLang"))
    if theme is None:
        theme = OxmlElement("w:themeFontLang")
        settings.append(theme)
    theme.set(qn("w:val"), "he-IL")
    theme.set(qn("w:bidi"), "he-IL")


def _configure(doc: Document) -> None:
    _set_doc_defaults_rtl(doc)
    _set_settings_rtl(doc)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_NAME
    normal.font.size = BODY_SIZE
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Cm(1.0)
        sec.bottom_margin = Cm(1.0)
        sec.left_margin = Cm(1.0)
        sec.right_margin = Cm(1.0)
        sp = sec._sectPr
        if sp.find(qn("w:bidi")) is None:
            sp.append(OxmlElement("w:bidi"))
        if sp.find(qn("w:pgBorders")) is None:
            pg = OxmlElement("w:pgBorders")
            pg.set(qn("w:offsetFrom"), "text")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "4")
                b.set(qn("w:color"), "000000")
                pg.append(b)
            sp.append(pg)


def _rule(doc, *, before=2, after=2) -> None:
    p = _para(doc, before=before, after=after)
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:color"), "000000")
    bdr.append(bot)
    p._p.get_or_add_pPr().append(bdr)


def _title(doc) -> None:
    _text(
        doc,
        "תחשיב זכויות אישי",
        bold=True,
        color=DARK_BLUE,
        size=TITLE_SIZE,
        after=6,
        before=1,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def _header(doc) -> None:
    _text(doc, "תאריך: {{TODAY}}", after=1, size=BODY_SIZE)
    _text(doc, "עבור: {{FULL_NAME}}", after=1, size=BODY_SIZE)
    _text(doc, "חשבון בנק: {{BANK_ACCOUNT}}", after=3, size=BODY_SIZE)
    _rule(doc, before=1, after=3)
    _text(
        doc,
        "להלן תחשיב הזכויות בהתאם לקריטריונים שאושרו בהחלטת האספה:",
        after=4,
        size=BODY_SIZE,
    )


def _calc_table(doc) -> None:
    p = _para(doc, after=8)
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.right_indent = Cm(0)
    r = p.add_run("{{p calc_table}}")
    _set_run_rtl(r)


def _section(doc, title: str, *, after=1, center=False, size=None) -> None:
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    _text(doc, title, bold=True, color=DARK_BLUE, size=size or SECTION_SIZE, after=after, align=align)


def _note(doc, flag, headline, body) -> None:
    _text(doc, f"{{%p if {flag} %}}", after=0, size=NOTES_SIZE)
    _text(doc, headline, bold=True, after=0, size=NOTES_SIZE)
    _text(doc, body, after=0, size=NOTES_SIZE)
    _text(doc, "{%p endif %}", after=0, size=NOTES_SIZE)


def _notes(doc) -> None:
    _text(doc, "{%p if show_NOTES_SECTION %}", after=0, size=NOTES_SIZE)
    _para(doc, before=10, after=2)
    _section(doc, "הערות והבהרות", after=2)
    _note(
        doc,
        "show_DEATH_SECTION",
        "תוספת בגין מקרה פטירה",
        "תוספת בגין מקרה פטירה: התוספת מתחלקת בין החברים הזכאים בהתאם לפירוט.",
    )
    _note(
        doc,
        "show_WORK_GRANT_SECTION",
        "קיזוז בגין מענק עידוד עבודה",
        "קיזוז בגין מענק עידוד עבודה כמפורט בתחשיב.",
    )
    _note(
        doc,
        "show_NEW_MEMBER_SECTION",
        "זכאותך תקום לאחר קבלתך לחברות בקיבוץ",
        "זכאותך תקום לאחר קבלתך לחברות בקיבוץ.",
    )
    _note(
        doc,
        "show_BUILDING_DEBT_SECTION",
        "חוב בגין בנייה פרטית",
        "ידוע לי כי יש לי חוב בגין בנייה פרטית אשר יקוזז מסכום זה.",
    )
    _text(doc, "{%p endif %}", after=0, size=NOTES_SIZE)


def _receipt(doc) -> None:
    _rule(doc, before=3, after=4)
    _section(doc, "כתב קבלה ושחרור", after=3, center=True, size=RECEIPT_TITLE_SIZE)
    _text(
        doc,
        'אני הח"מ {{FULL_NAME}}, לאחר שעיינתי בתחשיב הזכויות המפורט לעיל, מאשר/ת בזאת את נכונות החישוב ואת הסכומים המפורטים בו.',
        after=2,
        size=BODY_SIZE,
    )
    _text(
        doc,
        'אני מסכים/ה לקבלת הסכום הסופי בסך {{P}} לחשבון הבנק {{BANK_ACCOUNT}}, מאשר/ת קיזוז חובות כמפורט, ומוותר/ת על כל טענה עתידית.',
        after=4,
        size=BODY_SIZE,
    )
    _signature(doc)


def _date_space_paragraph(cell) -> None:
    """Reserve space for the PDF date field — no drawn box in DOCX."""
    p = cell.add_paragraph()
    _set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Cm(0.2)
    fmt.right_indent = Cm(0.2)

    p_pr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "560")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)

    _set_run_rtl(
        p.add_run(DATE_BOX_MARKER),
        size=Pt(1),
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )


def _signature_space_paragraph(cell) -> None:
    """Reserve vertical space for the PDF signature field — no drawn box in DOCX."""
    p = cell.add_paragraph()
    _set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT)
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Cm(4.6)
    fmt.right_indent = Cm(0.3)

    p_pr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "560")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)

    _set_run_rtl(
        p.add_run(SIG_BOX_MARKER),
        size=Pt(1),
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )


def _signature(doc) -> None:
    """Title + reserved space (left); date line (right). Box drawn only in PDF."""
    _para(doc, before=2, after=2)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _table_rtl(t)
    _table_width(t, 18.3)
    _row_height(t.rows[0], 900)

    date_c = t.rows[0].cells[0]
    sig_c = t.rows[0].cells[1]
    date_c.width = Cm(5.5)
    sig_c.width = Cm(12.8)

    _pad(date_c, t=40, b=90, s=50, e=30)
    date_c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    date_c.text = ""
    dp = date_c.paragraphs[0]
    _set_paragraph_rtl(dp, align=WD_ALIGN_PARAGRAPH.RIGHT)
    dp.paragraph_format.space_after = Pt(10)
    _set_run_rtl(dp.add_run("תאריך:"), size=BODY_SIZE)
    _date_space_paragraph(date_c)

    _pad(sig_c, t=40, b=90, s=20, e=40)
    sig_c.text = ""
    lp = sig_c.paragraphs[0]
    _set_paragraph_rtl(lp, align=WD_ALIGN_PARAGRAPH.CENTER)
    lp.paragraph_format.space_after = Pt(10)
    lp.paragraph_format.left_indent = Cm(4.6)
    lp.paragraph_format.right_indent = Cm(0.3)
    _set_run_rtl(lp.add_run("חתימה (שדה לחתימה דיגיטלית)"), size=Pt(9))
    _signature_space_paragraph(sig_c)
    _para(doc, after=3)


def _footer(doc) -> None:
    _rule(doc, before=3, after=3)
    _section(doc, "הנחיות להחזרה:", after=2, size=SECTION_SIZE)
    _text(doc, "1. ניתן להעביר את הטופס החתום במסירה ידנית למזכירות הקיבוץ.", after=1, size=FOOTER_SIZE)
    _text(doc, "2. בבארי: משרדי הקיבוץ / מזכירות הקיבוץ.", after=1, size=FOOTER_SIZE)
    p = _para(doc, after=2, line_spacing=1.2)
    parts = [
        ("3. ", {"bold": True}),
        ("את הטופס החתום – יש לשלוח במייל ל-", {"bold": True}),
        ("hakeren@beeri.co.il", {"color": LINK_BLUE, "underline": True, "bold": False}),
    ]
    for content, style in parts:
        r = p.add_run(content)
        _set_run_rtl(
            r,
            bold=style.get("bold", False),
            color=style.get("color"),
            underline=style.get("underline", False),
            size=FOOTER_SIZE,
        )
    _text(doc, "העברת הטופס החתום מהווה תנאי להמשך הטיפול בתשלום.", bold=True, after=1, size=FOOTER_SIZE)
    _text(doc, "במזכירות יתקבלו הטפסים החתומים בלבד.", bold=True, after=2, size=FOOTER_SIZE)
    _text(doc, "לשאלות ובירורים: 054-7918000", after=0, size=FOOTER_SIZE)


def build_with_python_docx(output_path: Path) -> Path:
    return create_fidelity_template_docx(output_path)


def create_fidelity_template_docx(output_path: Path) -> Path:
    doc = Document()
    _configure(doc)
    _title(doc)
    _header(doc)
    _calc_table(doc)
    _notes(doc)
    _receipt(doc)
    _footer(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def create_template(output_path: Path) -> Path:
    return build_with_python_docx(output_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "templates" / "תחשיב זכויות אישי.docx"
    print(f"Created: {build_with_python_docx(out)}")
