"""Build calculation table subdocument for docxtpl."""

from __future__ import annotations

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docxtpl import DocxTemplate

DARK_BLUE_HEX = "1F4E79"
TOTAL_ROW_HEX = "E8F0F8"
LAST_ROW_HEX = "FFF2CC"
WHITE_HEX = "FFFFFF"
FONT_NAME = "Arial"
TABLE_SIZE = Pt(10.5)
LINE_SPACING = 1.1
CELL_PAD_HEADER_PT = 7
CELL_PAD_NORMAL_PT = 6
CELL_PAD_TOTAL_PT = 8
CELL_PAD_SIDE_PT = 2
PAGE_WIDTH_CM = 21.0
PAGE_MARGIN_CM = 1.0
TABLE_WIDTH_CM = 14.9
COL_WIDTHS_CM = (1.3, 11.0, 2.6)


def _page_content_width_cm() -> float:
    return PAGE_WIDTH_CM - (2 * PAGE_MARGIN_CM)


def _jc_val(align) -> str:
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if align == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    return "right"


def _tbl_pr(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    return tbl_pr


def _table_rtl(table) -> None:
    tbl_pr = _tbl_pr(table)
    for tag in ("w:bidiVisual", "w:jc", "w:tblInd"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    tbl_pr.append(OxmlElement("w:bidiVisual"))
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    tbl_pr.append(jc)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)


def _table_width(table, cm: float) -> None:
    tbl_pr = _tbl_pr(table)
    for tag in ("w:tblW", "w:tblLayout"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(cm * 567)))
    w.set(qn("w:type"), "dxa")
    tbl_pr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _clear_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    for tag in ("w:tcBorders",):
        node = tc_pr.find(qn(tag))
        if node is not None:
            tc_pr.remove(node)
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tc_pr.append(b)


def _shade(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _borders(cell, color="BFBFBF") -> None:
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), color)
        b.append(e)
    cell._tc.get_or_add_tcPr().append(b)


def _set_paragraph_rtl(paragraph, *, align=WD_ALIGN_PARAGRAPH.RIGHT) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi", "w:textAlignment"):
        node = p_pr.find(qn(tag))
        if node is not None:
            p_pr.remove(node)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), _jc_val(align))
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))
    paragraph.alignment = align


def _set_run_ltr(run, *, bold=False, size=None) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or TABLE_SIZE


def _set_run_rtl(run, *, bold=False, size=None, white=False) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or TABLE_SIZE
    if white:
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:rtl", "w:lang"):
        node = r_pr.find(qn(tag))
        if node is not None:
            r_pr.remove(node)
    r_pr.append(OxmlElement("w:rtl"))
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    r_pr.append(lang)


def _pt_to_twips(pt: float) -> str:
    return str(int(round(pt * 20)))


def _set_line_spacing(paragraph, factor: float = LINE_SPACING) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    node = p_pr.find(qn("w:spacing"))
    if node is not None:
        p_pr.remove(node)
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:line"), str(int(round(factor * 240))))
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:before"), "0")
    sp.set(qn("w:after"), "0")
    p_pr.append(sp)


def _set_cell_padding(
    cell,
    *,
    top_pt: float = CELL_PAD_NORMAL_PT,
    bottom_pt: float = CELL_PAD_NORMAL_PT,
    side_pt: float = CELL_PAD_SIDE_PT,
) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:tcMar"))
    if node is not None:
        tc_pr.remove(node)
    mar = OxmlElement("w:tcMar")
    for side, pt in (
        ("top", top_pt),
        ("bottom", bottom_pt),
        ("left", side_pt),
        ("right", side_pt),
    ):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), _pt_to_twips(pt))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tc_pr.append(mar)


def _set_cell_valign(cell, align: str = "center") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:vAlign"))
    if node is not None:
        tc_pr.remove(node)
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), align)
    tc_pr.append(va)


def _cell_padding_for_row(*, header: bool = False, total: bool = False) -> tuple[float, float]:
    if header:
        return CELL_PAD_HEADER_PT, CELL_PAD_HEADER_PT
    if total:
        return CELL_PAD_TOTAL_PT, CELL_PAD_TOTAL_PT
    return CELL_PAD_NORMAL_PT, CELL_PAD_NORMAL_PT


def _cell(
    cell,
    content,
    *,
    bold=False,
    fill=None,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    white=False,
    amount=False,
    pad_top_pt: float | None = None,
    pad_bottom_pt: float | None = None,
):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = LINE_SPACING
    _set_line_spacing(p)
    text = content if not amount else f"\u200e{content}"
    r = p.add_run(text)
    if amount:
        _set_run_ltr(r, bold=bold, size=TABLE_SIZE)
    elif white:
        _set_run_rtl(r, bold=bold, size=TABLE_SIZE, white=True)
    else:
        _set_run_rtl(r, bold=bold, size=TABLE_SIZE)
    top = pad_top_pt if pad_top_pt is not None else CELL_PAD_NORMAL_PT
    bottom = pad_bottom_pt if pad_bottom_pt is not None else CELL_PAD_NORMAL_PT
    _set_cell_padding(cell, top_pt=top, bottom_pt=bottom)
    _set_cell_valign(cell)
    if fill:
        _shade(cell, fill)


def _row_fill(ri: int, row: dict, last_index: int) -> str:
    if ri == last_index:
        return LAST_ROW_HEX
    if row.get("is_total"):
        return TOTAL_ROW_HEX
    return WHITE_HEX


def _populate_calc_table(table, rows: list[dict]) -> None:
    _table_rtl(table)
    _table_width(table, TABLE_WIDTH_CM)

    for tr in table.rows:
        for i, w in enumerate(COL_WIDTHS_CM):
            tr.cells[i].width = Cm(w)

    headers = ("סעיף", "תיאור", "סכום (₪)")
    header_pad = _cell_padding_for_row(header=True)
    for i, h in enumerate(headers):
        _cell(
            table.rows[0].cells[i],
            h,
            bold=True,
            fill=DARK_BLUE_HEX,
            align=WD_ALIGN_PARAGRAPH.RIGHT,
            white=True,
            pad_top_pt=header_pad[0],
            pad_bottom_pt=header_pad[1],
        )
        _borders(table.rows[0].cells[i], color=DARK_BLUE_HEX)

    last_index = len(rows)
    for ri, row in enumerate(rows, 1):
        fill = _row_fill(ri, row, last_index)
        bold = bool(row.get("is_total")) or ri == last_index
        is_total = bool(row.get("is_total")) or ri == last_index
        pad_top, pad_bottom = _cell_padding_for_row(total=is_total)
        vals = (row["num"], row["desc"], row["amount"])
        for ci, val in enumerate(vals):
            _cell(
                table.rows[ri].cells[ci],
                val,
                bold=bold,
                fill=fill,
                align=WD_ALIGN_PARAGRAPH.RIGHT,
                amount=ci == 2,
                pad_top_pt=pad_top,
                pad_bottom_pt=pad_bottom,
            )
            _borders(table.rows[ri].cells[ci])


def build_calc_table_subdoc(template: DocxTemplate, rows: list[dict]):
    """Wrap data table in a full-width RTL shell so it sits flush right."""
    subdoc = template.new_subdoc()
    content_w = _page_content_width_cm()
    spacer_w = max(0.0, content_w - TABLE_WIDTH_CM)

    outer = subdoc.add_table(rows=1, cols=2)
    outer.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _table_rtl(outer)
    _table_width(outer, content_w)

    # RTL bidiVisual: cells[0] is the right column (data), cells[1] is left spacer.
    data_cell = outer.rows[0].cells[0]
    spacer_cell = outer.rows[0].cells[1]
    data_cell.width = Cm(TABLE_WIDTH_CM)
    spacer_cell.width = Cm(spacer_w)
    _clear_borders(data_cell)
    _clear_borders(spacer_cell)
    data_cell.text = ""
    spacer_cell.text = ""

    inner = data_cell.add_table(rows=1 + len(rows), cols=3)
    _populate_calc_table(inner, rows)
    return subdoc
