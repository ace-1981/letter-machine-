"""
Build DOCX template — single page, strict RTL, matched to reference image.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
DARK_BLUE_HEX = "1F4E79"
LIGHT_BLUE_HEX = "D9EAF7"
SIG_BOX_MARKER = "⟦SIG_BOX⟧"
DATE_BOX_MARKER = "⟦DATE_BOX⟧"
ZEBRA_HEX = "F2F2F2"
WHITE_HEX = "FFFFFF"
TOTAL_ROW_HEX = "E8F0F8"
LAST_ROW_HEX = "FFF2CC"
CALC_TABLE_WIDTH_CM = 14.9
CALC_COL_WIDTHS_CM = (1.7, 10.6, 2.6)
CALC_RIGHT_GAP_CM = 0.5
PAGE_WIDTH_CM = 21.0
PAGE_MARGIN_CM = 0.9
LINK_BLUE = RGBColor(0x05, 0x63, 0xC1)

FONT_NAME = "Arial"
BODY_SIZE = Pt(10.5)
TABLE_SIZE = Pt(10.5)
TITLE_SIZE = Pt(20)
SECTION_SIZE = Pt(11)
RECEIPT_TITLE_SIZE = Pt(12)
NOTES_SIZE = Pt(9.5)
FOOTER_SIZE = Pt(9.5)
TABLE_ROW_HEIGHT = 320
TABLE_HEADER_HEIGHT = 360
LINE_SPACING = 1.05

STYLE_BODY = "LG Body"
STYLE_SECTION = "LG Section"
STYLE_NOTES = "LG Notes"
STYLE_RECEIPT = "LG Receipt"
STYLE_FOOTER = "LG Footer"
STYLE_TITLE = "LG Title"


def _jc_val(align, *, mirror_for_bidi: bool = False) -> str:
    if align == WD_ALIGN_PARAGRAPH.CENTER:
        return "center"
    if mirror_for_bidi:
        if align == WD_ALIGN_PARAGRAPH.LEFT:
            return "right"
        return "left"
    if align == WD_ALIGN_PARAGRAPH.LEFT:
        return "left"
    return "right"


def _set_paragraph_rtl(
    paragraph,
    *,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    mirror_align: bool = False,
) -> None:
    # Use PHYSICAL alignment (jc=right) rather than the legacy "left"+bidi
    # mirror trick. Physical alignment survives Word edits even if Word strips
    # the paragraph's w:bidi, keeping the document right-aligned. bidi is still
    # applied for correct RTL reading order of mixed text.
    p_pr = paragraph._p.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi", "w:textAlignment"):
        node = p_pr.find(qn(tag))
        if node is not None:
            p_pr.remove(node)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), _jc_val(align, mirror_for_bidi=mirror_align))
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))
    fmt = paragraph.paragraph_format
    fmt.left_indent = Cm(0)
    fmt.right_indent = Cm(0)
    fmt.first_line_indent = Cm(0)


def _set_run_rtl(run, *, bold=False, size=None, color=None, underline=False) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or BODY_SIZE
    run.font.underline = underline
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:rtl", "w:lang"):
        node = r_pr.find(qn(tag))
        if node is not None:
            r_pr.remove(node)
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    r_pr.append(OxmlElement("w:rtl"))
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    r_pr.append(lang)
    if color:
        run.font.color.rgb = color


def _set_run_ltr(run, *, bold=False, size=None, color=None, underline=False) -> None:
    run.bold = bold
    run.font.name = FONT_NAME
    run.font.size = size or BODY_SIZE
    run.font.underline = underline
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:rtl", "w:lang"):
        node = r_pr.find(qn(tag))
        if node is not None:
            r_pr.remove(node)
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "en-US")
    lang.set(qn("w:bidi"), "en-US")
    r_pr.append(lang)
    if color:
        run.font.color.rgb = color


def _para(
    doc,
    *,
    after=0,
    before=0,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    line_spacing=LINE_SPACING,
    style: str | None = None,
):
    p = doc.add_paragraph(style=style)
    _set_paragraph_rtl(p, align=align)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.line_spacing = line_spacing
    return p


def _text(
    doc,
    content,
    *,
    bold=False,
    color=None,
    size=None,
    after=0,
    before=0,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    underline=False,
    style: str | None = None,
):
    p = _para(doc, after=after, before=before, align=align, style=style)
    r = p.add_run(content)
    _set_run_rtl(r, bold=bold, color=color, size=size, underline=underline)
    return p


def _mixed(doc, parts, *, after=1):
    p = _para(doc, after=after)
    for content, style in parts:
        r = p.add_run(content)
        _set_run_rtl(r, bold=style.get("bold", False), color=style.get("color"), underline=style.get("underline", False))
    return p


def _shade(cell, fill_hex: str) -> None:
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shd)


def _borders(cell, color="BFBFBF", dashed=False, sz="4") -> None:
    val = "dashed" if dashed else "single"
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), sz)
        e.set(qn("w:color"), color)
        b.append(e)
    cell._tc.get_or_add_tcPr().append(b)


def _clear_row_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:tcBorders"))
    if node is not None:
        tc_pr.remove(node)
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tc_pr.append(b)


def _pad(cell, t=55, b=55, s=70, e=70) -> None:
    m = OxmlElement("w:tcMar")
    for name, v in (("top", t), ("bottom", b), ("start", s), ("end", e)):
        n = OxmlElement(f"w:{name}")
        n.set(qn("w:w"), str(v))
        n.set(qn("w:type"), "dxa")
        m.append(n)
    cell._tc.get_or_add_tcPr().append(m)


def _row_height(row, height: int) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    for tag in ("w:trHeight",):
        node = tr_pr.find(qn(tag))
        if node is not None:
            tr_pr.remove(node)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(height))
    h.set(qn("w:hRule"), "atLeast")
    tr_pr.append(h)


def _set_cell_valign(cell, align: str = "center") -> None:
    va = OxmlElement("w:vAlign")
    va.set(qn("w:val"), align)
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:vAlign"))
    if node is not None:
        tc_pr.remove(node)
    tc_pr.append(va)


def _ensure_editable_settings(doc: Document) -> None:
    """Remove Word protection / read-only flags from document settings."""
    settings = doc.part.settings.element
    for tag in ("w:writeProtection", "w:documentProtection", "w:readOnlyRecommended"):
        node = settings.find(qn(tag))
        if node is not None:
            settings.remove(node)


def clear_docx_readonly(path: Path) -> None:
    """Clear OS read-only flag (e.g. after copy or ZIP extract on Windows)."""
    import os
    import stat

    if not path.is_file():
        return
    mode = path.stat().st_mode
    path.chmod(mode | stat.S_IWRITE)
    if os.name == "nt" and hasattr(path, "is_file"):
        try:
            import ctypes

            attrs = ctypes.windll.kernel32.GetFileAttributesW(str(path))
            if attrs != -1 and (attrs & 0x1):
                ctypes.windll.kernel32.SetFileAttributesW(str(path), attrs & ~0x1)
        except Exception:
            pass


def _cell(
    cell,
    content,
    *,
    bold=False,
    fill=None,
    align=WD_ALIGN_PARAGRAPH.RIGHT,
    white=False,
    size=None,
    amount=False,
):
    cell.text = ""
    p = cell.paragraphs[0]
    _set_paragraph_rtl(p, align=align, mirror_align=False)
    text = content if not amount else f"\u200e{content}"
    r = p.add_run(text)
    if amount:
        _set_run_ltr(r, bold=bold, size=size or TABLE_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF) if white else None)
    elif white:
        _set_run_rtl(r, bold=bold, size=size or TABLE_SIZE, color=RGBColor(0xFF, 0xFF, 0xFF))
    else:
        _set_run_rtl(r, bold=bold, size=size or TABLE_SIZE)
    _pad(cell, t=40 if not fill else 44, b=40 if not fill else 44)
    _set_cell_valign(cell)
    if fill:
        _shade(cell, fill)


def _page_content_width_cm() -> float:
    return PAGE_WIDTH_CM - (2 * PAGE_MARGIN_CM)


def _clear_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:tcBorders"))
    if node is not None:
        tc_pr.remove(node)
    b = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "nil")
        b.append(e)
    tc_pr.append(b)


def _tbl_pr(table):
    tbl_pr = table._tbl.tblPr
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        table._tbl.insert(0, tbl_pr)
    return tbl_pr


def _table_rtl(table) -> None:
    tbl_pr = _tbl_pr(table)
    for tag in ("w:bidiVisual", "w:jc", "w:tblInd"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    tbl_pr.append(OxmlElement("w:bidiVisual"))
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    tbl_pr.append(jc)
    ind = OxmlElement("w:tblInd")
    ind.set(qn("w:w"), "0")
    ind.set(qn("w:type"), "dxa")
    tbl_pr.append(ind)


def _table_width(table, cm: float) -> None:
    tbl_pr = _tbl_pr(table)
    for tag in ("w:tblW", "w:tblLayout"):
        node = tbl_pr.find(qn(tag))
        if node is not None:
            tbl_pr.remove(node)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), str(int(cm * 567)))
    w.set(qn("w:type"), "dxa")
    tbl_pr.append(w)
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)


def _set_style_rtl(style, *, align=WD_ALIGN_PARAGRAPH.RIGHT, size=None, bold: bool | None = None) -> None:
    style.font.name = FONT_NAME
    if size is not None:
        style.font.size = size
    if bold is not None:
        style.font.bold = bold
    style.paragraph_format.line_spacing = LINE_SPACING
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(0)

    p_pr = style.element.get_or_add_pPr()
    for tag in ("w:jc", "w:bidi"):
        node = p_pr.find(qn(tag))
        if node is not None:
            p_pr.remove(node)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), _jc_val(align, mirror_for_bidi=False))
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))

    r_pr = style.element.get_or_add_rPr()
    for tag in ("w:rtl", "w:lang", "w:rFonts"):
        node = r_pr.find(qn(tag))
        if node is not None:
            r_pr.remove(node)
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT_NAME)
    r_pr.append(r_fonts)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    r_pr.append(lang)
    r_pr.append(OxmlElement("w:rtl"))


def _add_paragraph_styles(doc: Document) -> None:
    from docx.enum.style import WD_STYLE_TYPE

    normal = doc.styles["Normal"]
    _set_style_rtl(normal, align=WD_ALIGN_PARAGRAPH.RIGHT, size=BODY_SIZE)

    defs = (
        (STYLE_BODY, WD_ALIGN_PARAGRAPH.RIGHT, BODY_SIZE, False),
        (STYLE_SECTION, WD_ALIGN_PARAGRAPH.RIGHT, SECTION_SIZE, True),
        (STYLE_NOTES, WD_ALIGN_PARAGRAPH.RIGHT, NOTES_SIZE, False),
        (STYLE_RECEIPT, WD_ALIGN_PARAGRAPH.RIGHT, BODY_SIZE, False),
        (STYLE_FOOTER, WD_ALIGN_PARAGRAPH.RIGHT, FOOTER_SIZE, False),
        (STYLE_TITLE, WD_ALIGN_PARAGRAPH.CENTER, TITLE_SIZE, True),
    )
    for name, align, size, bold in defs:
        if name in [s.name for s in doc.styles]:
            style = doc.styles[name]
        else:
            style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = normal
        _set_style_rtl(style, align=align, size=size, bold=bold)


def _set_doc_defaults_rtl(doc: Document) -> None:
    styles = doc.styles.element
    doc_defaults = styles.find(qn("w:docDefaults"))
    if doc_defaults is None:
        doc_defaults = OxmlElement("w:docDefaults")
        styles.insert(0, doc_defaults)

    p_pr_default = doc_defaults.find(qn("w:pPrDefault"))
    if p_pr_default is not None:
        doc_defaults.remove(p_pr_default)
    p_pr_default = OxmlElement("w:pPrDefault")
    p_pr = OxmlElement("w:pPr")
    jc = OxmlElement("w:jc")
    # Physical right alignment as the document-wide default: any paragraph whose
    # direct/style jc is stripped by a Word edit still inherits right alignment.
    jc.set(qn("w:val"), "right")
    p_pr.append(jc)
    p_pr.append(OxmlElement("w:bidi"))
    p_pr_default.append(p_pr)
    doc_defaults.append(p_pr_default)

    r_pr_default = doc_defaults.find(qn("w:rPrDefault"))
    if r_pr_default is not None:
        doc_defaults.remove(r_pr_default)
    r_pr_default = OxmlElement("w:rPrDefault")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT_NAME)
    r_fonts.set(qn("w:hAnsi"), FONT_NAME)
    r_fonts.set(qn("w:cs"), FONT_NAME)
    r_pr.append(r_fonts)
    lang = OxmlElement("w:lang")
    lang.set(qn("w:val"), "he-IL")
    lang.set(qn("w:bidi"), "he-IL")
    r_pr.append(lang)
    r_pr.append(OxmlElement("w:rtl"))
    r_pr_default.append(r_pr)
    doc_defaults.append(r_pr_default)


def _set_settings_rtl(doc: Document) -> None:
    settings = doc.part.settings.element
    theme = settings.find(qn("w:themeFontLang"))
    if theme is None:
        theme = OxmlElement("w:themeFontLang")
        settings.append(theme)
    theme.set(qn("w:val"), "he-IL")
    theme.set(qn("w:bidi"), "he-IL")


def _configure(doc: Document) -> None:
    _set_doc_defaults_rtl(doc)
    _set_settings_rtl(doc)
    _add_paragraph_styles(doc)

    for sec in doc.sections:
        sec.page_height = Cm(29.7)
        sec.page_width = Cm(21.0)
        sec.top_margin = Cm(PAGE_MARGIN_CM)
        sec.bottom_margin = Cm(PAGE_MARGIN_CM)
        sec.left_margin = Cm(PAGE_MARGIN_CM)
        sec.right_margin = Cm(PAGE_MARGIN_CM)
        sp = sec._sectPr
        if sp.find(qn("w:bidi")) is None:
            sp.append(OxmlElement("w:bidi"))
        if sp.find(qn("w:pgBorders")) is None:
            pg = OxmlElement("w:pgBorders")
            pg.set(qn("w:offsetFrom"), "text")
            for side in ("top", "left", "bottom", "right"):
                b = OxmlElement(f"w:{side}")
                b.set(qn("w:val"), "single")
                b.set(qn("w:sz"), "4")
                b.set(qn("w:space"), "4")
                b.set(qn("w:color"), "000000")
                pg.append(b)
            sp.append(pg)


def _rule(doc, *, before=1, after=1) -> None:
    p = _para(doc, before=before, after=after)
    bdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), "4")
    bot.set(qn("w:color"), "000000")
    bdr.append(bot)
    p._p.get_or_add_pPr().append(bdr)


def _title(doc) -> None:
    _text(
        doc,
        "תחשיב זכויות אישי",
        bold=True,
        color=DARK_BLUE,
        size=TITLE_SIZE,
        after=2,
        before=0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        style=STYLE_TITLE,
    )


def _header(doc) -> None:
    _text(doc, "תאריך: {{TODAY}}", after=0, size=BODY_SIZE, style=STYLE_BODY)
    _text(doc, "עבור: {{FULL_NAME}}", after=0, size=BODY_SIZE, style=STYLE_BODY)
    _text(doc, "חשבון בנק: {{BANK_ACCOUNT}}", after=2, size=BODY_SIZE, style=STYLE_BODY)
    _rule(doc, before=0, after=2)
    _text(
        doc,
        "להלן תחשיב הזכויות בהתאם לקריטריונים שאושרו בהחלטת האספה:",
        after=2,
        size=BODY_SIZE,
        style=STYLE_BODY,
    )


def _calc_table(doc) -> None:
    """Embedded calc table — editable in Word; RTL shell flushes table to the right."""
    _para(doc, after=1)
    content_w = _page_content_width_cm()
    left_spacer_w = max(0.0, content_w - CALC_TABLE_WIDTH_CM - CALC_RIGHT_GAP_CM)

    outer = doc.add_table(rows=1, cols=3)
    outer.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _table_rtl(outer)
    _table_width(outer, content_w)

    # RTL bidiVisual: cells[0]=rightmost. Keep a small gap off the right frame,
    # then the data table, then the remaining empty space on the left.
    right_gap_cell = outer.rows[0].cells[0]
    data_cell = outer.rows[0].cells[1]
    spacer_cell = outer.rows[0].cells[2]
    right_gap_cell.width = Cm(CALC_RIGHT_GAP_CM)
    data_cell.width = Cm(CALC_TABLE_WIDTH_CM)
    spacer_cell.width = Cm(left_spacer_w)
    _clear_borders(right_gap_cell)
    _clear_borders(data_cell)
    _clear_borders(spacer_cell)
    right_gap_cell.text = ""
    data_cell.text = ""
    spacer_cell.text = ""

    table = data_cell.add_table(rows=1, cols=3)
    _table_rtl(table)
    _table_width(table, CALC_TABLE_WIDTH_CM)

    headers = ("סעיף", "תיאור", "סכום (₪)")
    header_aligns = (
        WD_ALIGN_PARAGRAPH.CENTER,
        WD_ALIGN_PARAGRAPH.RIGHT,
        WD_ALIGN_PARAGRAPH.RIGHT,
    )
    for i, h in enumerate(headers):
        table.rows[0].cells[i].width = Cm(CALC_COL_WIDTHS_CM[i])
        _cell(
            table.rows[0].cells[i],
            h,
            bold=True,
            fill=DARK_BLUE_HEX,
            white=True,
            align=header_aligns[i],
        )
        _borders(table.rows[0].cells[i], color=DARK_BLUE_HEX)
    _row_height(table.rows[0], TABLE_HEADER_HEIGHT)

    def add_row(
        num: str,
        desc: str,
        amount: str,
        *,
        bold: bool = False,
        fill: str | None = None,
        row_if: str | None = None,
        static_amount: bool = False,
    ) -> None:
        if row_if:
            open_row = table.add_row()
            for i, w in enumerate(CALC_COL_WIDTHS_CM):
                open_row.cells[i].width = Cm(w)
            p_open = open_row.cells[0].paragraphs[0]
            _set_paragraph_rtl(p_open, mirror_align=False)
            p_open.text = f"{{%tr if {row_if} %}}"
            for ci in range(3):
                _clear_row_borders(open_row.cells[ci])

        row = table.add_row()
        for i, w in enumerate(CALC_COL_WIDTHS_CM):
            row.cells[i].width = Cm(w)
        _cell(row.cells[0], num, bold=bold, fill=fill, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell(row.cells[1], desc, bold=bold, fill=fill)
        _cell(
            row.cells[2],
            amount,
            bold=bold,
            fill=fill,
            amount=static_amount or amount.startswith("{{"),
        )
        for ci in range(3):
            _borders(row.cells[ci])
        _row_height(row, TABLE_ROW_HEIGHT)

        if row_if:
            close_row = table.add_row()
            for i, w in enumerate(CALC_COL_WIDTHS_CM):
                close_row.cells[i].width = Cm(w)
            p_close = close_row.cells[0].paragraphs[0]
            _set_paragraph_rtl(p_close, mirror_align=False)
            p_close.text = "{%tr endif %}"
            for ci in range(3):
                _clear_row_borders(close_row.cells[ci])

    add_row("1", "מספר שנים לצורך התחשיב", "{{ H }}")
    add_row("2", "סכום בגין שנת ותק אחת", "10 ₪", static_amount=True)
    add_row(
        "3",
        'סה"כ בגין ותק (מספר שנים כפול סכום לשנת ותק)',
        "{{ J }}",
        bold=True,
        fill=TOTAL_ROW_HEX,
    )
    add_row("4", "תוספת שווה לכל זכאי", "{{ I }}")
    add_row(
        "5",
        "תוספת בגין פטירה לכל זכאי",
        "{{ L }}",
        row_if="show_DEATH_SECTION",
    )
    add_row(
        "6",
        'סה"כ זכאות',
        "{{ M }}",
        bold=True,
        fill=TOTAL_ROW_HEX,
        row_if="show_WORK_GRANT_SECTION",
    )
    add_row(
        "6",
        'סה"כ זכאות',
        "{{ M }}",
        bold=True,
        fill=LAST_ROW_HEX,
        row_if="not show_WORK_GRANT_SECTION",
    )
    add_row(
        "7",
        "קיזוז בגין מענק עידוד עבודה",
        "{{ O }}",
        row_if="show_WORK_GRANT_SECTION",
    )
    add_row(
        "ח",
        "סה\"כ אחרי קיזוז 'מענק עידוד עבודה'",
        "{{ P }}",
        bold=True,
        fill=LAST_ROW_HEX,
        row_if="show_WORK_GRANT_SECTION",
    )
    _para(doc, after=2)


def _section(doc, title: str, *, after=0, center=False, size=None) -> None:
    align = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT
    _text(
        doc,
        title,
        bold=True,
        color=DARK_BLUE,
        size=size or SECTION_SIZE,
        after=after,
        align=align,
        style=STYLE_SECTION,
    )


def _note(doc, flag, headline, body) -> None:
    _text(doc, f"{{%p if {flag} %}}", after=0, size=NOTES_SIZE, style=STYLE_NOTES)
    _text(doc, headline, bold=True, after=0, size=NOTES_SIZE, style=STYLE_NOTES)
    _text(doc, body, after=0, size=NOTES_SIZE, style=STYLE_NOTES)
    _text(doc, "{%p endif %}", after=0, size=NOTES_SIZE, style=STYLE_NOTES)


def _notes(doc) -> None:
    _text(doc, "{%p if show_NOTES_SECTION %}", after=0, size=NOTES_SIZE, style=STYLE_NOTES)
    _para(doc, before=2, after=1)
    _section(doc, "הערות והבהרות", after=1)
    _note(
        doc,
        "show_DEATH_SECTION",
        "תוספת בגין מקרה פטירה",
        "תוספת בגין מקרה פטירה: התוספת מתחלקת בין החברים הזכאים בהתאם לפירוט.",
    )
    _note(
        doc,
        "show_WORK_GRANT_SECTION",
        "קיזוז בגין מענק עידוד עבודה",
        "קיזוז בגין מענק עידוד עבודה כמפורט בתחשיב.",
    )
    _note(
        doc,
        "show_NEW_MEMBER_SECTION",
        "זכאותך תקום לאחר קבלתך לחברות בקיבוץ",
        "זכאותך תקום לאחר קבלתך לחברות בקיבוץ.",
    )
    _note(
        doc,
        "show_BUILDING_DEBT_SECTION",
        "חוב בגין בנייה פרטית",
        "ידוע לי כי יש לי חוב בגין בנייה פרטית אשר יקוזז מסכום זה.",
    )
    _text(doc, "{%p endif %}", after=0, size=NOTES_SIZE, style=STYLE_NOTES)


def _receipt(doc) -> None:
    _rule(doc, before=2, after=2)
    _section(doc, "כתב קבלה ושחרור", after=2, center=True, size=RECEIPT_TITLE_SIZE)
    _text(
        doc,
        'אני הח"מ {{FULL_NAME}}, לאחר שעיינתי בתחשיב הזכויות המפורט לעיל, מאשר/ת בזאת את נכונות החישוב ואת הסכומים המפורטים בו.',
        after=1,
        size=BODY_SIZE,
        style=STYLE_RECEIPT,
    )
    _text(
        doc,
        'אני מסכים/ה לקבלת הסכום הסופי בסך {{P}} לחשבון הבנק {{BANK_ACCOUNT}}, מאשר/ת קיזוז חובות כמפורט, ומוותר/ת על כל טענה עתידית.',
        after=2,
        size=BODY_SIZE,
        style=STYLE_RECEIPT,
    )
    _signature(doc)


def _date_space_paragraph(cell) -> None:
    """Reserve space for the PDF date field — no drawn box in DOCX."""
    p = cell.add_paragraph()
    _set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, mirror_align=False)
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Cm(0.2)
    fmt.right_indent = Cm(0.2)

    p_pr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "560")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)

    _set_run_rtl(
        p.add_run(DATE_BOX_MARKER),
        size=Pt(1),
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )


def _signature_space_paragraph(cell) -> None:
    """Reserve vertical space for the PDF signature field — no drawn box in DOCX."""
    p = cell.add_paragraph()
    _set_paragraph_rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT, mirror_align=False)
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.left_indent = Cm(0.3)
    fmt.right_indent = Cm(0.3)

    p_pr = p._p.get_or_add_pPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "560")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)

    _set_run_rtl(
        p.add_run(SIG_BOX_MARKER),
        size=Pt(1),
        color=RGBColor(0xFF, 0xFF, 0xFF),
    )


def _signature(doc) -> None:
    """Title + reserved space (left); date line (right). Box drawn only in PDF."""
    _para(doc, before=1, after=4)
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.RIGHT
    _table_rtl(t)
    _table_width(t, 18.3)
    _row_height(t.rows[0], 720)

    date_c = t.rows[0].cells[0]
    sig_c = t.rows[0].cells[1]
    date_c.width = Cm(5.5)
    sig_c.width = Cm(12.8)

    _pad(date_c, t=20, b=50, s=50, e=30)
    date_c.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    date_c.text = ""
    dp = date_c.paragraphs[0]
    _set_paragraph_rtl(dp, align=WD_ALIGN_PARAGRAPH.RIGHT, mirror_align=False)
    dp.paragraph_format.space_after = Pt(4)
    _set_run_rtl(dp.add_run("תאריך:"), size=BODY_SIZE)
    _date_space_paragraph(date_c)

    _pad(sig_c, t=20, b=50, s=30, e=40)
    sig_c.text = ""
    lp = sig_c.paragraphs[0]
    _set_paragraph_rtl(lp, align=WD_ALIGN_PARAGRAPH.RIGHT, mirror_align=False)
    lp.paragraph_format.space_after = Pt(4)
    lp.paragraph_format.left_indent = Cm(0.3)
    lp.paragraph_format.right_indent = Cm(0.3)
    _set_run_rtl(lp.add_run("חתימה (שדה לחתימה דיגיטלית)"), size=Pt(9))
    _signature_space_paragraph(sig_c)
    _para(doc, after=12)


def _footer(doc) -> None:
    _rule(doc, before=10, after=2)
    _section(doc, "הנחיות להחזרה:", after=1, size=SECTION_SIZE)
    _text(doc, "1. ניתן להעביר את הטופס החתום במסירה ידנית למזכירות הקיבוץ.", after=0, size=FOOTER_SIZE, style=STYLE_FOOTER)
    _text(doc, "2. בבארי: משרדי הקיבוץ / מזכירות הקיבוץ.", after=0, size=FOOTER_SIZE, style=STYLE_FOOTER)
    p = _para(doc, after=1, line_spacing=LINE_SPACING, style=STYLE_FOOTER)
    parts = [
        ("3. ", {"bold": True}),
        ("את הטופס החתום – יש לשלוח במייל ל-", {"bold": True}),
        ("hakeren@beeri.co.il", {"color": LINK_BLUE, "underline": True, "bold": False}),
    ]
    for content, style in parts:
        r = p.add_run(content)
        _set_run_rtl(
            r,
            bold=style.get("bold", False),
            color=style.get("color"),
            underline=style.get("underline", False),
            size=FOOTER_SIZE,
        )
    _text(doc, "העברת הטופס החתום מהווה תנאי להמשך הטיפול בתשלום.", bold=True, after=0, size=FOOTER_SIZE, style=STYLE_FOOTER)
    _text(doc, "במזכירות יתקבלו הטפסים החתומים בלבד.", bold=True, after=1, size=FOOTER_SIZE, style=STYLE_FOOTER)
    _text(doc, "לשאלות ובירורים: 054-7918000", after=0, size=FOOTER_SIZE, style=STYLE_FOOTER)


def build_with_python_docx(output_path: Path) -> Path:
    return create_fidelity_template_docx(output_path)


def create_fidelity_template_docx(output_path: Path) -> Path:
    doc = Document()
    _configure(doc)
    _title(doc)
    _header(doc)
    _calc_table(doc)
    _notes(doc)
    _receipt(doc)
    _footer(doc)
    _ensure_editable_settings(doc)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    clear_docx_readonly(output_path)
    return output_path


def create_template(output_path: Path) -> Path:
    return build_with_python_docx(output_path)


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "templates" / "תחשיב זכויות אישי.docx"
    print(f"Created: {build_with_python_docx(out)}")
