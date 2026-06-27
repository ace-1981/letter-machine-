"""Proof: DOCX template edits flow to PDF without rebuild."""

from __future__ import annotations

import re
import shutil
import sys
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

TEMPLATE = ROOT / "templates" / "תחשיב זכויות אישי.docx"
CONFIG = ROOT / "templates" / "תחשיב זכויות אישי.json"
SAMPLE = ROOT / "samples" / "sample_data.xlsx"
OUT = ROOT / "samples" / "editable_template_proof"

ORIGINAL_TOTAL = 'סה"כ זכאות'
TEST_TOTAL = 'סה"כ זכאות TEST'
MARKER = "ROW_SPACING_PROOF_MARKER"


def _iter_tables(doc):
    for table in doc.tables:
        yield table
        for row in table.rows:
            for cell in row.cells:
                for nested in cell.tables:
                    yield nested


def _find_calc_table(doc):
    for table in _iter_tables(doc):
        if len(table.columns) != 3:
            continue
        if any(
            "מספר שנים לצורך התחשיב" in (cell.text or "")
            for row in table.rows
            for cell in row.cells
        ):
            return table
    return None


def _replace_document_xml(path: Path, replacer) -> None:
    tmp = path.with_suffix(".docx.tmp")
    with zipfile.ZipFile(path, "r") as zin, zipfile.ZipFile(tmp, "w") as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                text = data.decode("utf-8")
                text = replacer(text)
                data = text.encode("utf-8")
            zout.writestr(item, data)
    tmp.replace(path)


def _pdf_text(pdf_path: Path) -> str:
    import fitz

    doc = fitz.open(str(pdf_path))
    try:
        return "\n".join(page.get_text() for page in doc)
    finally:
        doc.close()


def _preview_pdf(excel: Path | None = None) -> Path:
    from src.letter_generator import generate_single_letter

    OUT.mkdir(parents=True, exist_ok=True)
    result = generate_single_letter(
        excel or SAMPLE,
        CONFIG,
        OUT,
        output_format="pdf",
        pdf_preferred="word",
    )
    return Path(result["pdf"])


def proof_text_edit() -> bool:
    backup = TEMPLATE.with_suffix(".docx.bak")
    shutil.copy2(TEMPLATE, backup)
    try:
        _replace_document_xml(
            TEMPLATE,
            lambda xml: xml.replace(ORIGINAL_TOTAL, TEST_TOTAL),
        )
        text = _pdf_text(_preview_pdf())
        compact = text.replace(" ", "")
        return "זכאותTEST" in compact or TEST_TOTAL in text
    finally:
        shutil.copy2(backup, TEMPLATE)
        backup.unlink(missing_ok=True)


def proof_spacing_edit() -> bool:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    backup = TEMPLATE.with_suffix(".docx.bak")
    shutil.copy2(TEMPLATE, backup)
    try:
        doc = Document(str(TEMPLATE))
        calc_table = _find_calc_table(doc)
        if calc_table is None or len(calc_table.rows) < 2:
            return False

        def set_first_data_row_padding(top: int, bottom: int) -> None:
            row = calc_table.rows[1]
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                node = tc_pr.find(qn("w:tcMar"))
                if node is not None:
                    tc_pr.remove(node)
                mar = OxmlElement("w:tcMar")
                for side, val in (("top", top), ("bottom", bottom), ("start", 70), ("end", 70)):
                    el = OxmlElement(f"w:{side}")
                    el.set(qn("w:w"), str(val))
                    el.set(qn("w:type"), "dxa")
                    mar.append(el)
                tc_pr.append(mar)

        def read_row_padding(row_index: int) -> int | None:
            row = calc_table.rows[row_index]
            tc_pr = row.cells[0]._tc.tcPr
            if tc_pr is None:
                return None
            mar = tc_pr.find(qn("w:tcMar"))
            if mar is None:
                return None
            top = mar.find(qn("w:top"))
            return int(top.get(qn("w:w"))) if top is not None else None

        set_first_data_row_padding(120, 120)
        doc.save(str(TEMPLATE))

        from src.letter_generator import generate_single_letter

        test_excel = ROOT / "samples" / "editable_template_proof" / "test_one.xlsx"
        if not test_excel.is_file():
            import pandas as pd

            test_excel.parent.mkdir(parents=True, exist_ok=True)
            row = {chr(65 + j): "" for j in range(20)}
            row.update({
                "C": 50001, "E": "כהן", "F": "ישראל",
                "H": 10, "I": 1000, "J": 100, "L": 500, "M": 1600,
                "O": 200, "P": 1400, "S": "12-345-678901", "R": "", "T": "",
            })
            pd.DataFrame([row]).to_excel(test_excel, index=False)

        OUT.mkdir(parents=True, exist_ok=True)
        edited_docx = Path(
            generate_single_letter(
                test_excel, CONFIG, OUT, output_format="docx"
            )["docx"]
        )
        edited_pad = None
        out_doc = Document(str(edited_docx))
        out_calc = _find_calc_table(out_doc)
        if out_calc is not None:
            tc_pr = out_calc.rows[1].cells[0]._tc.tcPr
            if tc_pr is not None:
                mar = tc_pr.find(qn("w:tcMar"))
                if mar is not None:
                    top = mar.find(qn("w:top"))
                    edited_pad = int(top.get(qn("w:w"))) if top is not None else None

        pdf_big = _preview_pdf(test_excel)
        shutil.copy2(backup, TEMPLATE)
        baseline_docx = Path(
            generate_single_letter(
                test_excel, CONFIG, OUT, output_format="docx"
            )["docx"]
        )
        pdf_normal = _preview_pdf(test_excel)

        import fitz

        def row2_y(pdf: Path) -> float:
            page = fitz.open(str(pdf))[0]
            hits = page.search_for("סכום בגין שנת ותק")
            return min(r.y0 for r in hits) if hits else 0.0

        y_big = row2_y(pdf_big)
        y_norm = row2_y(pdf_normal)
        print(
            f"spacing padding top={edited_pad} (expected 120); "
            f"PDF row2 Y edited={y_big:.1f} baseline={y_norm:.1f}"
        )
        return edited_pad == 120
    finally:
        if backup.is_file():
            shutil.copy2(backup, TEMPLATE)
            backup.unlink(missing_ok=True)


if __name__ == "__main__":
    ok1 = proof_text_edit()
    ok2 = proof_spacing_edit()
    print("text_edit:", ok1)
    print("spacing_edit:", ok2)
    raise SystemExit(0 if ok1 and ok2 else 1)
